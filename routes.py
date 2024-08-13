from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import render_template, request, redirect, url_for, jsonify, g
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from gensim.models import Phrases
from gensim.models.phrases import Phraser
from project import app as applet
from nltk import sent_tokenize as st
from nltk import word_tokenize as wt
from pymongo import MongoClient
from collections import defaultdict
import sqlite3, re, os, docx, pymystem3, chardet, docx, time, pickle, gensim, math, io, base64
from prereform2modern import Processor
morph = pymystem3.Mystem(weight=False,grammar_info=False)
morph = pymystem3.Mystem(weight=False)
jwt = JWTManager(applet)
userData = sqlite3.connect('userdata.db')
cur = userData.cursor()
cur.execute('CREATE TABLE IF NOT EXISTS users(russian TEXT, login TEXT, password TEXT)')
userData.commit()

mongo = MongoClient()
CURRENT_DIR = os.getcwd()

morph = pymystem3.Mystem(weight=False)

class normalizer:
    def word_modernizer(word):
        preresult = Processor.process_text(word,show=False,delimiters=False,check_brackets=False)
        result = preresult[0]
        return result

    def new_gettext(file):
        try:
            ##print(f'{file}')
            if file.endswith('.txt'):
                text = open(rf'{file}','rb')
                text_body = text.read()
                enc = chardet.detect(text_body).get("encoding")
                ##print(enc)
                if enc and enc.lower() != "utf-8" and enc.lower() != "windows-1251":
                    text_body = text_body.decode(enc)
                    text_body = text_body.encode("utf-8")
                    text_body = text_body.decode("utf-8")
                    ##print('Открыт текст!')
                    return text_body
                elif enc and enc.lower() == "windows-1251":
                    text = open(rf'{file}', 'r', encoding = 'windows-1251')
                    text_body = text.read()
                    text.close()
                    ##print('Открыт текст!')
                    return text_body
                else:
                    text = open(rf'{file}', 'r', encoding = 'UTF-8')
                    text_body = text.read()
                    text.close()
                    ##print('Открыт текст!')
                    return text_body
            elif file.endswith('.docx'):
                doc = docx.Document(rf'{file}')
                text = (paragraph.text for paragraph in doc.paragraphs)
                text_body = '\n'.join(text)
                return text_body
            else:
                ##print('Неподдерживаемый формат!')
                pass
        except:
            pass 

    def insert_breaks(text):
                if text == None:
                    text = ''
                else:
                    text = text
                text = re.sub('([А-Я]).([А-Я]). ([А-Я])',r'\1.\2._\3', text) #Here we try to save names like В.И. Ленин by underscoring spaces in them
                text = re.sub('([А-Я]). ([А-Я]). ([А-Я])',r'\1._\2._\3', text) #Here we try to save names like Л. Д. Троцкий by underscoring spaces in them
                sents = st(text)
                with_breaks = ' BREAK '.join(sents)
                with_breaks = re.sub('_',' ', with_breaks) #Here we replace our underscores by regular spaces
                with_breaks = with_breaks + ' BREAK'
                ##print('Разделители вставлены!')
                return with_breaks

    def linearize(result):
        if 'analysis' not in result:
            text = result['text']
            lex = 'NONE'
            gr = 'NONE'
        else:
            if len(result['analysis']) == 0:
                text = result['text']
                lex = 'NONE'
                gr = 'NONE'
            else:
                text = result['text']
                analysis = result['analysis'][0]
                lex = analysis['lex']
                gr = [item for item in re.split('[,=|()]',analysis['gr']) if item != '']
        if text == ' ' or text=='\n':
            pass
        else:
            return {'token': text, 'lemma':lex, 'grammar': gr}
        
    def nolemmatize(text):
        tokenized = wt(text)
        result = [{'token':tok,'lemma':tok,'grammar':'X'} for tok in tokenized]
        return result

        
    def separator(jsn):
        sents = defaultdict(list)
        count = 0
        token_count = 0
        sents[0] = []
        for item in jsn:
            if item['token'] != 'BREAK':
                item['number'] = token_count
                token_count += 1
                sents[count].append(item)
            else:
                count += 1
                sents[count] = []
                token_count = 0
        sents = {num: sents[num] for num in sents if sents[num] != []}
        return sents
    
def preprocess(file,destination,depth,user,dbname): 
    try:
        text = normalizer.new_gettext(file)
        ##print(text)
        ##print(f'Извлечен текст из файла {file}')
        text_breaks = normalizer.insert_breaks(text)
        raw_sents = re.sub('\n',' ',text_breaks)
        raw_sents = re.sub('&#1123;','ѣ',raw_sents)
        raw_sents = re.sub('&#1122;','Ѣ',raw_sents)
        raw_sents = re.sub('&#1139;','ѳ',raw_sents)
        raw_sents = re.sub('&#1138;','Ѳ',raw_sents)
        raw_sents = re.sub('^.',' ',raw_sents)
        raw_sents = re.split('BREAK',raw_sents)[:-1]
        raw_sents = dict(enumerate(raw_sents))
        for_pm3 = re.sub('''[*.,:;?!()"«»“„_—>']''','', text_breaks)
        for_pm3 = for_pm3.replace('[','').replace(']','')
        for_pm3 = re.sub('\n',' ', for_pm3)
        if depth == 'fullscale':
            analyzed = morph.analyze(for_pm3)
            analyzed = [item for item in [normalizer.linearize(entry) for entry in analyzed] if item != None]
        elif depth == 'shallow':
            analyzed = normalizer.nolemmatize(for_pm3)
        numbered_sents = normalizer.separator(analyzed)
        #Здесь добавляю предложения в файл для векторизации
        prepared = [numbered_sents[num] for num in numbered_sents]
        def sent_to_destination(sent):
            sentwords = [f"{word['lemma']}_{word['grammar'][0]}" for word in sent if word != 'NONE']
            sentwords = ' '.join(sentwords) + '\n'
            destination.write(sentwords)
        [sent_to_destination(sent) for sent in prepared]
        #добавление кончается
        lengths = [len(numbered_sents[sent]) for sent in numbered_sents] #here we count the number of tokens in every sent
        cardinality = sum(lengths) #here we sum these lengths and get the cardinality of the file, i.e. how many tokens it contains
        ##print(raw_sents)
        ##print(numbered_sents)
        ##print(f'Текст {file} обработан!')
        ##print(file)
        ##print(numbered_sents)
        print(file)
        result = [{'sentnum': num, 'file': file, 'raw': raw_sents[num], 'numbered': numbered_sents[num]} for num in raw_sents]
    except:
        cardinality = 0
        result = [{'sentnum':0,'file':file, 'raw':'','numbered':[]}]
    
    db = mongo[f'{user}'] #Maybe it would make sense to create separate dbs for separate users
    collection = db[f'{dbname}'] #The same holds as to collections
    collection.insert_many(result)
    os.remove(file)

    return {'cardinality':cardinality}

@applet.route('/hello', methods = ['GET','POST'])
def hello():
    global CURRENT_DIR
    ##print(CURRENT_DIR)
    return 'Шалом!'

@applet.route('/uploader', methods = ['GET','POST'])
def uploader():
        global CURRENT_DIR
        directory = CURRENT_DIR
        CURRENT_USER = request.headers.get('login')
        #print(CURRENT_USER)
        if os.path.isdir(rf"{directory}/docx_files/{CURRENT_USER}"):
            ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf"{directory}/docx_files/{CURRENT_USER}")
            ##print('Директория создана!')
        files = request.files.getlist('file[]')
        #print('POSTFILES!')
        names = []
        #print('POSTNAMES!')
        for file in files:
            name = file.filename
            names.append(name)
            #print(name)
            file.save(rf"{directory}/docx_files/{CURRENT_USER}/{name}")
            #print('NAME_SAVED!')
            ##print(rf"{directory}/docx_files/{CURRENT_USER}/{name}")
            ##print(f'{name} saved!')
        return {'type':'namelist', 'namelist': names}

@applet.route('/concordance_saver', methods = ['GET','POST'])
def concord_saver():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    concordance = request.json.get('concordance')
    dbname = request.json.get('DBname')
    concname = request.json.get('ConcName')
    if os.path.isdir(rf"{directory}/concordances/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/concordances/{CURRENT_USER}")
        ##print('Директория создана!')
    all_info = {'database':dbname,'concordance':concordance}
    pickle.dump(all_info, open(rf"{directory}/concordances/{CURRENT_USER}/{dbname}_{concname}_concordance","wb"))
    return {'result':'success'}


@applet.route('/download_ngram',methods=['GET','POST'])
def downgrams():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    word_1 =  request.json.get('w_1')
    word_2 =  request.json.get('w_2')
    instances = request.json.get('instances')
    examples = request.json.get('result')
    dbname = request.json.get('DBname')
    concname = request.json.get('ConcName')
    if os.path.isdir(rf"{directory}/for_download/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/for_download/{CURRENT_USER}")
        ##print('Директория создана!')
    doc = docx.Document()
    if examples != 'NONE':
        parabody = doc.add_paragraph()
        parabody.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for example in examples:
            text = f"{example['frst']} {example['scnd']}\n"
            parabody.add_run(text)
    if examples == 'NONE' and word_1 != '' and word_2 != '':
        parainst = doc.add_paragraph()
        parainst.add_run(f'При скачивании просматривались следующие примеры сочетания\n {word_1.upper()} + {word_2.upper()}:\n\n')
        for instance in instances:
            txt = f"{instance['raw']} [{instance['file']}]"
            parainst.add_run(txt)
    doc.save(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")
    bt = open(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx",'rb')
    bt = bt.read()
    bytes = base64.b64encode(bt).decode('ascii')
    os.remove(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")

    return {'for_download': bytes, 'name':concname}

@applet.route('/download_collsearch',methods=['GET','POST'])
def downcolls():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    word_1 = request.json.get('word_1').upper()
    word_2 = request.json.get('word_2').upper()
    word_3 = request.json.get('word_3')
    word_4 = request.json.get('word_4')
    if word_3 == None:
        word_3 = 'None'
    else:
        word_3 = word_3.upper()
    if word_4 == None:
        word_4 = 'None'
    else:
        word_4 = word_4.upper()
    examples = request.json.get('result')
    howmuch = request.json.get('howmuch')
    dbname = request.json.get('DBname')
    concname = request.json.get('ConcName')
    if os.path.isdir(rf"{directory}/for_download/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/for_download/{CURRENT_USER}")
        ##print('Директория создана!')
    doc = docx.Document()
    parahead = doc.add_paragraph()
    if word_3 == 'None' and word_4 == 'None':
        parahead.add_run(f'Первое слово: {word_1}, второе слово: {word_2}\n')
    elif word_3 != 'None' and word_4 == 'None':
        parahead.add_run(f'Первое слово: {word_1}, второе слово: {word_2}, третье слово: {word_3}\n')
    else:
        parahead.add_run(f'Первое слово: {word_1}, второе слово: {word_2}, третье слово: {word_3}, четвертое слово: {word_4}\n')
    parahead.add_run(f'Абсолютная частота: {howmuch}\n')
    parabody = doc.add_paragraph()
    parabody.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for example in examples:
        text = f"{example['raw']}[{example['file']}]"
        parabody.add_run(text)
    doc.save(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")
    bt = open(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx",'rb')
    bt = bt.read()
    bytes = base64.b64encode(bt).decode('ascii')
    os.remove(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")

    return {'for_download': bytes, 'name':concname}

@applet.route('/download_wordsearch',methods=['GET','POST'])
def downwords():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    word = request.json.get('word')
    word = word.upper()
    examples = request.json.get('result')
    howmuch = request.json.get('howmuch')
    ipm = request.json.get('ipm')
    dbname = request.json.get('DBname')
    concname = request.json.get('ConcName')
    if os.path.isdir(rf"{directory}/for_download/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/for_download/{CURRENT_USER}")
        ##print('Директория создана!')
    doc = docx.Document()
    parahead = doc.add_paragraph()
    parahead.add_run(f'{word}\n')
    parahead.add_run(f'Абсолютная частота: {howmuch}\n')
    parahead.add_run(f'Единиц на миллион: {ipm}\n')
    parabody = doc.add_paragraph()
    parabody.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for example in examples:
        text = f"{example['raw']}[{example['file']}]"
        parabody.add_run(text)
    doc.save(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")
    bt = open(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx",'rb')
    bt = bt.read()
    bytes = base64.b64encode(bt).decode('ascii')
    os.remove(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")

    return {'for_download': bytes, 'name':concname}


@applet.route('/download_concordance',methods = ['GET','POST'])
def downconc():
    try:
        global CURRENT_DIR
        directory = CURRENT_DIR
        CURRENT_USER = request.json.get('login')
        concordance = request.json.get('concordance')
        dbname = request.json.get('DBname')
        concname = request.json.get('ConcName')
        if os.path.isdir(rf"{directory}/for_download/{CURRENT_USER}"):
            ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf"{directory}/for_download/{CURRENT_USER}")
            ##print('Директория создана!')
        doc = docx.Document()
        for item in concordance:
            parahead = doc.add_paragraph()
            header = item['header']
            freq = item['frequency']
            ipm = item['ipm']
            head = parahead.add_run(f'{header}\n')
            head.font.bold
            absfreq = parahead.add_run(f'Абсолютная частота: {freq}\n')
            absfreq.font.bold        
            absfreq.font.italic
            aipiem = parahead.add_run(f'Единиц на миллион: {ipm}\n')
            aipiem.font.bold
            aipiem.font.italic
            examples = item['contents']
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for example in examples:
                para.add_run(f"{example['raw']}[{example['file']}]")
            doc.add_paragraph()
        doc.save(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")
        bt = open(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx",'rb')
        bt = bt.read()
        bytes = base64.b64encode(bt).decode('ascii')
        os.remove(rf"{directory}/for_download/{CURRENT_USER}/{dbname}_{concname}.docx")
        return {'for_download': bytes, 'name':concname}
    except Exception as err:
        #print(err)
        return {'exception': err}

    

@applet.route('/enumerate_concordances', methods = ['GET','POST'])
def enumerator():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    searchdir = rf"{directory}/concordances/{CURRENT_USER}"
    result = [{'name':name} for name in os.listdir(searchdir)]
    return {'concordances':result}


@applet.route('/concordance_opener', methods = ['GET','POST'])
def concord_opener():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    concordance = request.json.get('ConcName')
    conc_jsn = pickle.load(open(rf"{directory}/concordances/{CURRENT_USER}/{concordance}","rb"))
    return {'result': conc_jsn}

@applet.route('/concordance_deleter',methods = ['GET','POST'])
def concord_del():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    concordance = request.json.get('ConcName')
    os.remove(rf"{directory}/concordances/{CURRENT_USER}/{concordance}")
    return {'deletion_alert':concordance}

@applet.route('/delete', methods = ['GET','POST'])
def deleter():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    if os.path.isdir(rf"{directory}/docx_files/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/docx_files/{CURRENT_USER}")
    directive = request.json.get('directive')
    if directive == 'remove':
        existent_files = os.listdir(rf'{directory}/docx_files/{CURRENT_USER}')
        for file in existent_files:
            os.remove(rf'{directory}/docx_files/{CURRENT_USER}/{file}')
    return {'type':'removed_warning','namelist':existent_files}

@applet.route('/file_list', methods = ['GET','POST'])
def give_filenames():
    global CURRENT_DIR
    directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    if os.path.isdir(rf"{directory}/docx_files/{CURRENT_USER}"):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf"{directory}/docx_files/{CURRENT_USER}")
    ##print(directory)
    directive = request.json.get('directive')
    names = []
    if directive == 'give_filenames':
        existent_files = os.listdir(rf'{directory}/docx_files/{CURRENT_USER}')
        for file in existent_files:
            names.append(file)
        return {'type':'files','namelist':names}


@applet.route('/clear', methods = ['GET','POST'])
def clear():
    CURRENT_USER = request.json.get('login')
    dbname = request.json.get('dbname')
    db = mongo[f'{CURRENT_USER}']
    collection = db[f'{dbname}']
    collection.drop()
    return {'deletion':f'База данных {dbname} удалена!'}

@applet.route('/dbinfo', methods = ['GET','POST'])
def dbtechinfo():
    CURRENT_USER = request.json.get('login')
    dbname = request.json.get('dbname')
    db = mongo[f'{CURRENT_USER}']
    collection = db[f'{dbname}']
    cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
    cardinality = cardinality['cardinality']
    files = collection.distinct('file')
    files = list(files)
    grammar = collection.distinct('numbered.grammar')
    grammar = list(grammar)
    ##print(grammar)
    if len(grammar)==1 and 'X' in grammar:
        gr = 'нет'
    else:
        gr = 'есть'
    return {'cardinality':cardinality,'files':files, 'grammar':gr}


@applet.route('/database', methods = ['GET','POST'])
def process():
    #current_directory = os.getcwd()
    dbname = request.json.get('dbname')
    depth = request.json.get('depth')
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    files = os.listdir(rf'{current_directory}/docx_files/{CURRENT_USER}')
    if os.path.isdir(rf'{current_directory}/for_vectors/{CURRENT_USER}'):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf'{current_directory}/for_vectors/{CURRENT_USER}')
        ##print('Директория создана!')
    ##print(files)
    for_vectors = open(rf'{current_directory}/for_vectors/{CURRENT_USER}/{dbname}.txt','w',encoding='UTF-8')
    jsons = [preprocess(f'{current_directory}/docx_files/{CURRENT_USER}/{file}',for_vectors,depth,CURRENT_USER,dbname) for file in files]
    cardinality = sum([file['cardinality'] for file in jsons])
    #jsons = [jsn for file in jsons for jsn in file['result'] if file['result'] != None]
    #jsons = [jsn for jsn in jsons]
    ##print(jsons)
    db = mongo[f'{CURRENT_USER}'] #Maybe it would make sense to create separate dbs for separate users
    collection = db[f'{dbname}'] #The same holds as to collections
    collection.insert_one({'cardinality':cardinality})
    print('123')
    #collection.insert_many(jsons)
    return {'type':'db_ready','namelist': files}

@applet.route('/append_database',methods=['GET','POST'])
def db_appender():
    #current_directory = os.getcwd()
    dbname = request.json.get('dbname')
    depth = request.json.get('depth')
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    files = os.listdir(rf'{current_directory}/docx_files/{CURRENT_USER}')
    if os.path.isdir(rf'{current_directory}/for_vectors/{CURRENT_USER}'):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf'{current_directory}/for_vectors/{CURRENT_USER}')
        ##print('Директория создана!')
    ##print(files)
    for_vectors = open(rf'{current_directory}/for_vectors/{CURRENT_USER}/{dbname}.txt','w',encoding='UTF-8')
    jsons = [preprocess(f'{current_directory}/docx_files/{CURRENT_USER}/{file}',for_vectors,depth,CURRENT_USER,dbname) for file in files]
    cardinality = sum([file['cardinality'] for file in jsons])
    #jsons = [jsn for file in jsons for jsn in file['result'] if file['result'] != None]
    #jsons = [jsn for jsn in jsons]
    ##print(jsons)
    db = mongo[f'{CURRENT_USER}'] #Maybe it would make sense to create separate dbs for separate users
    collection = db[f'{dbname}'] #The same holds as to collections
    old_cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
    collection.delete_one({'cardinality':{'$exists':True}})
    old_cardinality = old_cardinality['cardinality']
    cardinality = old_cardinality + cardinality
    collection.insert_one({'cardinality':cardinality})
    print('123')
    #collection.insert_many(jsons)
    return {'type':'db_ready','namelist': files}


@applet.route('/old_append_database',methods=['GET','POST'])
def append_db():
    #current_directory = os.getcwd()
    dbname = request.json.get('dbname')
    depth = request.json.get('depth')
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    ##print(current_directory)
    ##print(CURRENT_USER)
    files = os.listdir(rf'{current_directory}/docx_files/{CURRENT_USER}')
    if os.path.isdir(rf'{current_directory}/for_vectors/{CURRENT_USER}'):
        ##print('Такая директория уже существует!')
        pass
    else:
        os.mkdir(rf'{current_directory}/for_vectors/{CURRENT_USER}')
        ##print('Директория создана!')
    print(files)
    for_vectors = open(rf'{current_directory}/for_vectors/{CURRENT_USER}/{dbname}.txt','a',encoding='UTF-8')
    jsons = [preprocess(f'{current_directory}/docx_files/{CURRENT_USER}/{file}',for_vectors,depth,CURRENT_USER,dbname) for file in files]
    cardinality = sum([file['cardinality'] for file in jsons])
    #jsons = [jsn for file in jsons for jsn in file['result'] if file['result'] != None]
    #jsons = [jsn for jsn in jsons]
    ##print(jsons)
    db = mongo[f'{CURRENT_USER}'] #Maybe it would make sense to create separate dbs for separate users
    collection = db[f'{dbname}'] #The same holds as to collections
    collection.insert_one({'cardinality':cardinality})
    #collection.insert_many(jsons)
    return {'type':'db_ready','namelist': files}

@applet.route('/list_dbs', methods = ['GET','POST'])
def list_dbs():
    CURRENT_USER = request.json.get('login')
    db = mongo[f'{CURRENT_USER}']
    collections = db.list_collection_names()
    #return {'collection_names':collections}
    collections = [{'name':item} for item in collections]
    result = {'collection_names':collections}

    return result

@applet.route('/search', methods = ['GET','POST'])
def search():
    CURRENT_USER = request.json.get('login')
    ##print(CURRENT_USER)
    #word = input('Введите искомое слово: \n').lower().strip()
    try:
        word = request.json.get('word').lower().strip()
        searchtype = request.json.get('lemOrTok')
        dbname = request.json.get('DBname')
        ##print(dbname)
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        if searchtype == 'lemma':
            result = list(collection.find({'numbered.lemma': f'{word}'},{'_id':0,'numbered':0}))
        elif searchtype == 'token':
            result = list(collection.find({'numbered.token': f'{word}'},{'_id':0,'numbered':0}))
        elif searchtype == 'substring':
            result = list(collection.find({'numbered.token': {'$regex':f'^{word}+'}},{'_id':0,'numbered':0}))
        absolute = len(result) #number of entries
        ipm = (absolute/cardinality)*(10**6)
        return {'howmuch':absolute, 'ipm':format(ipm, '.2f'),'result':result, 'error':'0', 'cardinality':cardinality}
    except Exception as Err:
        return {'error':'1'}

@applet.route('/collocates', methods = ['GET','POST'])
def collocates():
    CURRENT_USER = request.json.get('login')
    try:
        word_1 = request.json.get('word_1').lower().strip()
        word_2 = request.json.get('word_2').lower().strip()
        dbname = request.json.get('DBname')
        lemtok_1 = request.json.get('lemOrTok_1')
        lemtok_2 = request.json.get('lemOrTok_2')
        ##print(word_1)
        ##print(word_2)
        ##print(lemtok_1)
        ##print(lemtok_2)
        mini_1 = int(request.json.get('mini_1'))
        maxi_1 = int(request.json.get('maxi_1'))
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        start = time.perf_counter()
        if lemtok_1 in ['lemma','token'] and lemtok_2 in ['lemma','token']:
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 == 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.token':{'$regex':f'^{word_2}+'}}]},{'_id':0}))
        elif lemtok_1 == 'substring' and lemtok_2 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.token':{'$regex':f'^{word_1}+'},f'numbered.{lemtok_2}':word_2}]},{'_id':0}))
        ##print('Предложения подобраны!')
        def fit_variants(sent):
            numd = sent['numbered'] #получаем для предложения список его токенов с характеристиками
            if lemtok_1 == 'lemma':
                candidates_1 = [entry['number'] for entry in numd if entry['lemma'] == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'token':
                candidates_1 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'substring':
                candidates_1 = [entry['number'] for entry in numd if word_1 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 1
            if lemtok_2 == 'lemma':
                candidates_2 = [entry['number'] for entry in numd if entry['lemma'] == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'token':
                candidates_2 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'substring':
                candidates_2 = [entry['number'] for entry in numd if word_2 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2

            detracts = [(j-i-1) for j in candidates_2 for i in candidates_1] #из каждого кандидата 2 вычитаем каждый кандидат 1 и ещё единицу
            for detract in detracts:
                if detract in range(mini_1,maxi_1+1):
                    return {'raw':sent['raw'],'file':sent['file'],'sentnum':sent['sentnum']}
            else:
                pass #иначе пропускаем, возвращая None
        result = [sent for sent in [fit_variants(sent) for sent in sents] if sent != None] #применяем fit_variants ко всем предложениям, потом фильтруем ненулевые
        absolute = len(result) #number of entries
        return {'howmuch':absolute, 'result':result,'error':'0','cardinality':cardinality}
    except Exception as err:
        return {'error':'1'}

@applet.route('/trigram_collocates', methods = ['GET','POST'])
def trigram_collocates():
    CURRENT_USER = request.json.get('login')
    try:
        word_1 = request.json.get('word_1').lower().strip()
        word_2 = request.json.get('word_2').lower().strip()
        word_3 = request.json.get('word_3').lower().strip()
        dbname = request.json.get('DBname')
        lemtok_1 = request.json.get('lemOrTok_1')
        lemtok_2 = request.json.get('lemOrTok_2')
        lemtok_3 = request.json.get('lemOrTok_3')
        mini_1 = int(request.json.get('mini_1'))
        maxi_1 = int(request.json.get('maxi_1'))
        mini_2 = int(request.json.get('mini_2'))
        maxi_2 = int(request.json.get('maxi_2'))
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        start = time.perf_counter()
        if lemtok_1 != 'substring' and lemtok_2 != 'substring' and lemtok_3 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2,f'numbered.{lemtok_3}':word_3}]},{'_id':0}))
        elif lemtok_1 == 'substring' and lemtok_2 != 'substring' and lemtok_3 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.token':{'$regex':f'^{word_1}+'},f'numbered.{lemtok_2}':word_2,f'numbered.{lemtok_3}':word_3}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 == 'substring' and lemtok_3 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.token':{'$regex':f'^{word_2}+'},f'numbered.{lemtok_3}':word_3}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 != 'substring' and lemtok_3 == 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2,f'numbered.token':{'$regex':f'^{word_3}+'}}]},{'_id':0}))
        ##print('Предложения подобраны!')
        def fit_variants(sent):
            numd = sent['numbered'] #получаем для предложения список его токенов с характеристиками
            if lemtok_1 == 'lemma':
                candidates_1 = [entry['number'] for entry in numd if entry['lemma'] == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'token':
                candidates_1 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'substring':
                candidates_1 = [entry['number'] for entry in numd if word_1 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 1            
            if lemtok_2 == 'lemma':
                candidates_2 = [entry['number'] for entry in numd if entry['lemma'] == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'token':
                candidates_2 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'substring':
                candidates_2 = [entry['number'] for entry in numd if word_2 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2            
            if lemtok_3 == 'lemma':
                candidates_3 = [entry['number'] for entry in numd if entry['lemma'] == word_3] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_3 == 'token':
                candidates_3 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_3] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_3 == 'substring':
                candidates_3 = [entry['number'] for entry in numd if word_3 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2            
            detracts_1_2 = set([(j-i-1) for j in candidates_2 for i in candidates_1]) #из каждого кандидата 2 вычитаем каждый кандидат 1 и ещё единицу
            detracts_2_3 = set([(p-q-1) for p in candidates_3 for q in candidates_2]) #из каждого кандидата 3 вычитаем каждый кандидат 2 и ещё единицу
            first_range = set(range(mini_1,maxi_1+1))
            second_range = set(range(mini_2,maxi_2+1))
            if len(first_range.intersection(detracts_1_2)) > 0 and len(second_range.intersection(detracts_2_3)) > 0:
                return {'raw':sent['raw'],'file':sent['file'],'sentnum':sent['sentnum']}
            else:
                pass

        result = [sent for sent in [fit_variants(sent) for sent in sents] if sent != None] #применяем fit_variants ко всем предложениям, потом фильтруем ненулевые
        finish = time.perf_counter() 
        absolute = len(result) #number of entries
        return {'howmuch':absolute, 'result':result, 'error': '0', 'cardinality':cardinality}
    except:
        return {'error':'1'}

@applet.route('/quadrigram_collocates', methods = ['GET','POST'])
def quadrigram_collocates():
    CURRENT_USER = request.json.get('login')
    try:
        word_1 = request.json.get('word_1').lower().strip()
        word_2 = request.json.get('word_2').lower().strip()
        word_3 = request.json.get('word_3').lower().strip()
        word_4 = request.json.get('word_4').lower().strip()
        dbname = request.json.get('DBname')
        lemtok_1 = request.json.get('lemOrTok_1')
        lemtok_2 = request.json.get('lemOrTok_2')
        lemtok_3 = request.json.get('lemOrTok_3')
        lemtok_4 = request.json.get('lemOrTok_4')
        ##print(word_1)
        ##print(word_2)
        ##print(word_3)
        ##print(word_4)
        ##print(lemtok_1)
        ##print(lemtok_2)
        ##print(lemtok_3)
        ##print(lemtok_4)
        mini_1 = int(request.json.get('mini_1'))
        maxi_1 = int(request.json.get('maxi_1'))
        mini_2 = int(request.json.get('mini_2'))
        maxi_2 = int(request.json.get('maxi_2'))
        mini_3 = int(request.json.get('mini_3'))
        maxi_3 = int(request.json.get('maxi_3'))
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        start = time.perf_counter()
        if lemtok_1 != 'substring' and lemtok_2 != 'substring' and lemtok_3 != 'substring' and lemtok_4 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2,f'numbered.{lemtok_3}':word_3,f'numbered.{lemtok_4}':word_4}]},{'_id':0}))
        elif lemtok_1 == 'substring' and lemtok_2 != 'substring' and lemtok_3 != 'substring' and lemtok_4 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.token':{'$regex':f'^{word_1}+'},f'numbered.{lemtok_2}':word_2,f'numbered.{lemtok_3}':word_3,f'numbered.{lemtok_4}':word_4}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 == 'substring' and lemtok_3 != 'substring' and lemtok_4 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.token':{'$regex':f'^{word_2}+'},f'numbered.{lemtok_3}':word_3,f'numbered.{lemtok_4}':word_4}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 != 'substring' and lemtok_3 == 'substring' and lemtok_4 != 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2,f'numbered.token':{'$regex':f'^{word_3}+'},f'numbered.{lemtok_4}':word_4}]},{'_id':0}))
        elif lemtok_1 != 'substring' and lemtok_2 != 'substring' and lemtok_3 != 'substring' and lemtok_4 == 'substring':
            sents = list(collection.find({'$and':[{f'numbered.{lemtok_1}':word_1,f'numbered.{lemtok_2}':word_2,f'numbered.{lemtok_3}':word_3,f'numbered.token':{'$regex':f'^{word_4}+'}}]},{'_id':0}))
        
        ##print('Предложения подобраны!')
        def fit_variants(sent):
            numd = sent['numbered'] #получаем для предложения список его токенов с характеристиками
            if lemtok_1 == 'lemma':
                candidates_1 = [entry['number'] for entry in numd if entry['lemma'] == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'token':
                candidates_1 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_1] #номера всех токенов, лемма которых равна слову 1
            elif lemtok_1 == 'substring':
                candidates_1 = [entry['number'] for entry in numd if word_1 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 1            
            if lemtok_2 == 'lemma':
                candidates_2 = [entry['number'] for entry in numd if entry['lemma'] == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'token':
                candidates_2 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_2] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_2 == 'substring':
                candidates_2 = [entry['number'] for entry in numd if word_2 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2            
            if lemtok_3 == 'lemma':
                candidates_3 = [entry['number'] for entry in numd if entry['lemma'] == word_3] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_3 == 'token':
                candidates_3 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_3] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_3 == 'substring':
                candidates_3 = [entry['number'] for entry in numd if word_3 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2            
            if lemtok_4 == 'lemma':
                candidates_4 = [entry['number'] for entry in numd if entry['lemma'] == word_4] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_4 == 'token':
                candidates_4 = [entry['number'] for entry in numd if normalizer.word_modernizer(entry['token'].lower().strip()) == word_4] #номера всех токенов, лемма которых равна слову 2
            elif lemtok_4 == 'substring':
                candidates_4 = [entry['number'] for entry in numd if word_4 in normalizer.word_modernizer(entry['token'].lower().strip())] #номера всех токенов, лемма которых равна слову 2            

            detracts_1_2 = set([(j-i-1) for j in candidates_2 for i in candidates_1]) #из каждого кандидата 2 вычитаем каждый кандидат 1 и ещё единицу
            detracts_2_3 = set([(p-q-1) for p in candidates_3 for q in candidates_2]) #из каждого кандидата 3 вычитаем каждый кандидат 2 и ещё единицу
            detracts_3_4 = set([(r-s-1) for r in candidates_4 for s in candidates_3]) #из каждого кандидата 4 вычитаем каждый кандидат 3 и ещё единицу
            first_range = set(range(mini_1,maxi_1+1))
            second_range = set(range(mini_2,maxi_2+1))
            third_range = set(range(mini_3,maxi_3+1))
            if len(first_range.intersection(detracts_1_2)) > 0 and len(second_range.intersection(detracts_2_3)) > 0 and len(third_range.intersection(detracts_3_4)) > 0:
                return {'raw':sent['raw'],'file':sent['file'],'sentnum':sent['sentnum']}
            else:
                pass

        result = [sent for sent in [fit_variants(sent) for sent in sents] if sent != None] #применяем fit_variants ко всем предложениям, потом фильтруем ненулевые
        finish = time.perf_counter()
        absolute = len(result) #number of entries
        return {'howmuch':absolute, 'result':result, 'error':'0', 'cardinality':cardinality}
    except:
        return {'error':'1'}

@applet.route('/full_concordance',methods=['GET','POST'])
def full_concordance():
    CURRENT_USER = request.json.get('login')
    try:
        upper = request.json.get('upper')
        lower = request.json.get('lower')
        dbname = request.json.get('DBname')
        pos = request.json.get('pos')
        ##print(pos)
        conctype = request.json.get('conctype')
        if upper == None:
            upper = 'А'
        if lower == None:
            lower = 'ЯЯЯЯ'
        upper = upper.upper()
        lower = lower.upper()
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        #Here we except the entry containing the key 'cardinality', which does not contain other keys
        first = list(collection.find({'cardinality':{'$not':{'$exists':True}}},{'sentnum':1,'raw':1,'file':1,'numbered.lemma':1,'numbered.grammar':1}))
        ##print('Список составлен!')
        lemmata = defaultdict(list)
        for item in first:
            if pos == 'all':            
                numbered = item['numbered']
                raw = item['raw']
                file = item['file']
                sentnum = item['sentnum']
                ##print(item)
                for entry in numbered:
                    lemma = entry['lemma'].upper()
                    ##print(f'Обрабатываю лемму {lemma}')
                    if lemma != 'NONE':
                        if lemma not in lemmata:
                            lemmata[lemma] = [{'sentnum':sentnum,'raw':raw, 'file':file}]
                        else:
                            lemmata[lemma].append({'sentnum':sentnum,'raw':raw,'file':file})
                    else:
                        pass
            else:
                numbered = item['numbered']
                raw = item['raw']
                file = item['file']
                sentnum = item['sentnum']
                ##print(item)
                for entry in numbered:
                    lemma = entry['lemma'].upper()
                    ##print(f'Обрабатываю лемму {lemma}')
                    if lemma != 'NONE' and entry['grammar'][0]==pos:
                        if lemma not in lemmata:
                            lemmata[lemma] = [{'sentnum':sentnum,'raw':raw, 'file':file}]
                        else:
                            lemmata[lemma].append({'sentnum':sentnum,'raw':raw,'file':file})
                    else:
                        pass


        lemmas = [{'header':elt, 'contents':lemmata[elt], 'frequency':len(lemmata[elt]),
                    'ipm':format((len(lemmata[elt])/cardinality)*10**6, '.2f')} for elt in lemmata]
        if conctype == 'alphabetic':
            sorted_lemmas = sorted(lemmas, key=lambda x: x['header'])
        elif conctype == 'reversed':
            sorted_lemmas = sorted(lemmas, key=lambda x: x['header'][::-1])
        elif conctype == 'frequency':
            sorted_lemmas = sorted(lemmas, key=lambda x: 1/(x['frequency']))
        chosen_lemmas = [item for item in sorted_lemmas if (item['header']>=upper and item['header']<=lower)]
        return {'content':chosen_lemmas,'error':'0','dbname':dbname,'cardinality':cardinality}
    except Exception as err:
        return {'error':'1'}
    
@applet.route('/broader_context',methods=['GET','POST'])
def broaden():
    CURRENT_USER = request.json.get('login')
    file = request.json.get('filename')
    sentnum = request.json.get('sentence')
    sentnum = int(sentnum)
    dbname = request.json.get('DBname')
    ##print(file,sentnum,dbname)
    db = mongo[f'{CURRENT_USER}']
    collection = db[f'{dbname}']
    cand_rng = [sentnum-i for i in range(0,10)] #Результат вычитания из номера предложения чисел до 9
    pos_rng = [x for x in cand_rng if x >= 0] #Фильтруем все неотрицательные
    min_val = min(pos_rng) #Находим среди них наименьшее, это будет левая граница
    rng = [num for num in range(min_val,sentnum+11)]
    context = collection.find({'sentnum':{'$in':rng}},{'_id':0,'raw':1,'file':1}) #вытаскиваем все предложения с номерами в нужном диапазоне
    sents = [item['raw'] for item in list(context) if item['file']==file]
    sents = ' '.join(sents)
    ##print(sents)
    return {'broadened':sents}

@applet.route('/register',methods=['GET','POST'])
def create_account():
    russian = request.json.get('russian')
    login = request.json.get('nomen')
    password = request.json.get('tessera')
    user_data = sqlite3.connect('userdata.db')
    cursor = user_data.cursor()
    cursor.execute('INSERT INTO users(russian,login,password) VALUES(?,?,?)',[russian,login,password])
    user_data.commit()
    return {'msg': f'Аккаунт на имя {russian} успешно создан'}

@applet.route('/log_in', methods=['GET','POST'])
def logging_in():
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    master_login = 'miserere_nobis'
    master_password = 'vanitas_vanitatum'
    usdata = sqlite3.connect('userdata.db')
    curs = usdata.cursor()
    logpass = list(curs.execute('SELECT login,password FROM users'))
    login = request.json.get('nomen')
    CURRENT_USER = login
    ##print(CURRENT_USER)
    password = request.json.get('tessera')
    if login == master_login and password == master_password:
        if os.path.isdir(rf'{current_directory}/for_vectors/{CURRENT_USER}'):
        ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf'{current_directory}/for_vectors/{CURRENT_USER}')
        ##print(f'Логин: {login}, пароль: {password}.')
        #access_token = create_access_token(identity=login)
        #return jsonify(access_token=access_token,login=login,status='success')
        return {'login': login,'status':'success'}
    elif (login,password) in logpass:
        if os.path.isdir(rf'{current_directory}/for_vectors/{CURRENT_USER}'):
        ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf'{current_directory}/for_vectors/{CURRENT_USER}')
        ##print(f'Логин: {login}, пароль: {password}.')
        #access_token = create_access_token(identity=login)
        return {'login': login,'status':'success'}
    else:
        #return jsonify(access_token='invalid',login=login, status='aborted')
        return jsonify(login=login, status='aborted')
    
@applet.route('/log_out',methods=['GET','POST'])
def logging_out():
    return jsonify(access_token='loggedout',login='выход',status='loggedout')

#Here begins the referencer part of the program

@applet.route('/referencer', methods=['GET','POST'])
def referencer():
    global CURRENT_DIR
    ##print('Ты прошёл!')
    database = sqlite3.connect(rf'{CURRENT_DIR}/actual_references.db')
    cur = database.cursor()
    query = request.json.get('word')
    strict = request.json.get('strict')
    #auth = request.json.get('Authorization')
    ##print(auth)
    if query != None:
        query = query.upper().strip()
    else:
        query = ''

    result = list(cur.execute(f'''SELECT *
                        FROM refers
                        '''))
    #[##print(item[1]) for item in result]
    
    if strict == 'yes':
        necessary = [quadruple for quadruple in result if re.sub(' ','',quadruple[3])==query]
    elif strict == 'no':
        necessary = [quadruple for quadruple in result if re.sub(' ','',quadruple[3]).startswith(query)]


    res_json = [{'headword':re.sub('\.','',quadruple[0]),'example':quadruple[2], 'issue':quadruple[4], 'link':quadruple[3]} for quadruple in necessary]

    dense_json = defaultdict(list)

    for item in res_json:
        headword = item['headword']
        the_rest = {'example':item['example'],'issue':item['issue'],'link':item['link']}
        if headword not in dense_json:
            dense_json[headword] = [the_rest]
        else:
            dense_json[headword].append(the_rest)


    output = [{'type':'reference','headword':word, 'issue':dense_json[word][0]['issue'], 'body':dense_json[word]} for word in dense_json]
    return output

@applet.route('/substrings', methods=['GET','POST'])
def substrings():
    global CURRENT_DIR
    database = sqlite3.connect(rf'{CURRENT_DIR}/actual_references.db')
    cur = database.cursor()
    query = request.json.get('word')
    strict = request.json.get('strict')
    if query != None:
        query = query.lower().strip()
    else:
        query = ''
    result = cur.execute(f'''SELECT *
                        FROM for_substrings
                        ''')
    result = list(result)
    result = [{'headword':item[0],'tokens':wt(item[1]),'example':item[2],'issue':item[3]} for item in result]
    ##print(result)

    
    def issubstring(word, goal):
        if goal.startswith(word):
            return 'TRUE'
        else:
            return 'FALSE'

    #if strict == 'yes':
    #    necessary = [quadruple for quadruple in result if ('◊' in quadruple[2] or '~' in quadruple[2]) and quadruple[3]=='NONE' and quadruple[1] == query]
    #elif strict == 'no':
    #    necessary = [quadruple for quadruple in result if ('◊' in quadruple[2] or '~' in quadruple[2]) and quadruple[3]=='NONE' and quadruple[1].startswith(query)]
    
    if strict == 'yes':
        query=morph.lemmatize(query)
        query=query[0]
        necessary = [quadruple for quadruple in result if query in quadruple['tokens']]
    elif strict == 'no':
        necessary = [quadruple for quadruple in result if 'TRUE' in [issubstring(query,token) for token in quadruple['tokens']]]
    res_json = [{'headword':quadruple['headword'],'example':quadruple['example'], 'issue':quadruple['issue'], 'link':'no'} for quadruple in necessary]
    
    #res_json = [{'headword':'ТУПОЙ', 'example':'БУРАТИНО БЫЛ ТУПОЙ','issue':'ПЕСЕНКА ПСОЯ','link':'нифига'}]
    dense_json = defaultdict(list)

    for item in res_json:
        headword = item['headword']
        the_rest = {'example':item['example'],'issue':item['issue'],'link':item['link']}
        if headword not in dense_json:
            dense_json[headword] = [the_rest]
        else:
            if the_rest not in dense_json[headword]:
                dense_json[headword].append(the_rest)
            else:
                pass

    output = [{'type':'substrings','headword':word, 'issue':dense_json[word][0]['issue'], 'body':dense_json[word]} for word in dense_json]
    return output

@applet.route('/refupload',methods=['GET','POST'])
def uploading():
        global CURRENT_DIR
        directory = CURRENT_DIR
        CURRENT_USER = request.files.get('login')
        ##print(CURRENT_USER)
        if os.path.isdir(rf"{directory}/referencer_docx"):
            ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf"{directory}/referencer_docx")
            ##print('Директория создана!')
        files = request.files.getlist('file[]')
        ##print(files)
        names = []
        for file in files:
            name = file.filename
            names.append(name)
            file.save(rf"{directory}/referencer_docx/{name}")
            ##print(f'{name} saved!')
        return {'type':'namelist', 'namelist': names}

@applet.route('/reffilelist',methods=['GET','POST'])
def filename_giver():
    global CURRENT_DIR
    directory = CURRENT_DIR
    ##print(directory)
    directive = request.json.get('directive')
    names = []
    if directive == 'give_filenames':
        existent_files = os.listdir(rf'{directory}/referencer_docx')
        for file in existent_files:
            names.append(file)
        return {'type':'files','namelist':names}

@applet.route('/refdelete', methods=['GET','POST'])
def deletion():
    global CURRENT_DIR
    directory = CURRENT_DIR
    directive = request.json.get('directive')
    if directive == 'remove':
        existent_files = os.listdir(rf'{directory}/referencer_docx')
        for file in existent_files:
            os.remove(rf'{directory}/referencer_docx/{file}')
    return {'type':'removed_warning','namelist':existent_files}

@applet.route('/remdb',methods=['GET','POST'])
def remove_db():
        global CURRENT_DIR
        database = sqlite3.connect(rf'{CURRENT_DIR}/actual_references.db')
        cur = database.cursor()
        cur.execute('PRAGMA encoding="utf-8"')
        cur.execute('CREATE TABLE IF NOT EXISTS refers(headword TEXT, head TEXT, example TEXT, link TEXT, issue TEXT)')
        cur.execute('DELETE FROM refers')
        cur.execute('CREATE TABLE IF NOT EXISTS for_substrings(headword TEXT, lemmata TEXT, example TEXT, issue TEXT)')
        cur.execute('DELETE FROM for_substrings')
        database.commit()

        return {'status':'removed'}

@applet.route('/refdb', methods=['GET','POST'])
def create_refdb():
    directive = request.json.get('directive')
    if directive == 'make_db':
        global CURRENT_DIR
        #os.remove('actual_references.db')
        database = sqlite3.connect(rf'{CURRENT_DIR}/actual_references.db')
        cur = database.cursor()
        cur.execute('PRAGMA encoding="utf-8"')
        cur.execute('CREATE TABLE IF NOT EXISTS refers(headword TEXT, head TEXT, example TEXT, link TEXT, issue TEXT)')
        #cur.execute('DELETE FROM refers')
        cur.execute('CREATE TABLE IF NOT EXISTS for_substrings(headword TEXT, lemmata TEXT, example TEXT, issue TEXT)')
        #cur.execute('DELETE FROM for_substrings')
        res = ''

        current_directory = os.getcwd()
        ##print(current_directory)
        files = os.listdir(rf'{current_directory}/referencer_docx')
        ##print(files)

    def process(file):
        ##print(file)
        doc = docx.Document(rf'{current_directory}/referencer_docx/{file}')
        lastpara = doc.add_paragraph() #ДОБАВЛЯЕМ ПОСЛЕДНИЙ ЛИПОВЫЙ ПАРАГРАФ, ЧТОБЫ ПРИ ПЕРЕХОДЕ К НЕМУ ОБРАБОТАЛСЯ ПРЕДПОСЛЕДНИЙ -- ПОСЛЕДНИЙ НАСТОЯЩИЙ.
        lastpara.add_run('END') #В ЛИПОВЫЙ ПАРАГРАФ ДОБАВЛЯЕМ ЛИПОВЫЙ ЗАГОЛОВОК, ЧТОБЫ ЗАВЕРШИТЬ ЦИКЛ ДОБАВЛЕНИЯ В ПРЕДЫДУЩУЮ СТАТЬЮ. ИЗБЕГАЕМ СИМВОЛОВ РИМСКИХ ЦИФР.
        issue = file
        entries = []
        entry = {}
        result = ''
        heading = ''

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                font = run.font
                if font.all_caps:
                    run.text = run.text.upper()
                else:
                    run.text = run.text
            
            if '@' not in paragraph.text and '#' not in paragraph.text:
                raw_text = paragraph.text.strip(' ')
                raw_text = re.sub('^[◁ ◀]','',raw_text)
                raw_text = re.sub('^[0-9]+\.','', raw_text)
                raw_text = re.sub('[<>]','',raw_text)
                raw_text = re.sub('^—','',raw_text)
                raw_text = re.sub('◊','Rhomb',raw_text) #Попробуем заменить ромбики на условные последовательности
                raw_text = re.sub('~','Tilde',raw_text) #Попробуем заменить ромбики на условные последовательности
                raw_text = raw_text.replace("[","sqr_brkt_lft")                
                raw_text = raw_text.replace("[","sqr_brkt_lft")
                raw_text = raw_text.replace("]","sqr_brkt_rht")
                raw_text = raw_text.replace("]","sqr_brkt_rht")
                raw_text = raw_text.replace("(","rnd_brkt_lft")
                raw_text = raw_text.replace(")","rnd_brkt_rht")
                raw_text = raw_text.strip()

                # ПРОБУЕМ ОТЛОВИТЬ ЖИРНЫЕ ПОСЛЕ ТИЛЬДЫ
                if raw_text.startswith('Tilde'):
                        pairs = []
                        for this_run in paragraph.runs:
                            this_font = this_run.font
                            if this_font.bold and '~' not in this_run.text and this_run.text[0].isupper():
                                alphabet = {'с', 'а', 'ш', 'з', 'ф', 'ч', 'ж', 'б', 'р', 'д', 'ы', 'ь', 'т', 'ю', 'л', 'п', 'й', 'э', 'е', 'я', 'м', 'н', 'щ', 'ъ', 'о', 'г', 'х', 'ё', 'к', 'в', 'и', 'ц', 'у'}
                                setrun = set(this_run.text)
                                if len(setrun.intersection(alphabet)) > 0:
                                    was = this_run.text
                                    was = was.replace("[","sqr_brkt_lft")                
                                    was = was.replace("[","sqr_brkt_lft")
                                    was = was.replace("]","sqr_brkt_rht")
                                    was = was.replace("]","sqr_brkt_rht")
                                    was = was.replace("(","rnd_brkt_lft")
                                    was = was.replace(")","rnd_brkt_rht")
                                    ##print(was)
                                    will = f'Tilde {was}' 
                                    ##print(will)
                                    pair = {'was':was,'will':will}
                                    pairs.append(pair)
                        ##print(pairs)
                        for pair in pairs:
                            raw_text = re.sub(rf'{pair["was"]}',rf'{pair["will"]}',raw_text) 
                raw_text = re.sub('Tilde Tilde','Tilde',raw_text)
                raw_text = raw_text.replace("sqr_brkt_lft","[")
                raw_text = raw_text.replace("sqr_brkt_rht","]")
                raw_text = raw_text.replace("rnd_brkt_lft","(")
                raw_text = raw_text.replace("rnd_brkt_rht",")")
                # КОНЕЦ ОТЛОВА ЖИРНЫХ ПОСЛЕ ТИЛЬДЫ
                while raw_text != '' and not raw_text[0].isalnum():
                    raw_text = raw_text[1:]
                initword = raw_text.split(' ', 1)[0] #выделяем заголовочное слово по первому пробелу
                if 'Х.' in initword:
                    initword = re.sub('Х.','X.', initword)
                elif 'З.' in initword:
                    initword = re.sub('З.','3.', initword)
                else:
                    initword = initword    
                initword = initword.replace('A','А')       
                initword = initword.replace('K','К')
                initword = initword.replace('O','О')
                initword = initword.replace('E','Е')
                initword = initword.replace('T','Т')
                initword = initword.replace('P','Р')
                initword = initword.replace('H','Н')
                initword = initword.replace('C','С')
                initword = initword.replace('B','В')
                initword = initword.replace('M','М')
                first_word = initword.strip('.,') #убираем у заголовочного слова точки и запятые
                ##print(first_word)
                #ТО ЕСТЬ ДО ЭТОГО МОМЕНТА ВЫЯСНЯЕМ, ЕСТЬ ЛИ В ПАРАГРАФЕ ЗАГОЛОВОК. ОК.

                forbidden = {'I','V','X','.','~'}

                condition_1 = (not first_word.isupper() and paragraph.text != '') or (first_word.isupper() and len(set(first_word).intersection(forbidden))>0)
                subcondition = initword == f'{first_word},' or initword == f'{first_word}.'
                condition_2 = first_word.isupper() and (len(first_word) > 1 or (len(first_word) == 1 and subcondition)) and len(paragraph.text) != 0 and len(set(first_word).intersection(forbidden)) == 0                
                if condition_1: #если верно условие 1, то есть первое слово не заголовочное
                    result = result + '\n' + raw_text
                elif condition_2: #если же слово заголовочное, то:
                    #здесь мы обрабатываем статью, попавшую перед нашим заголовком; 
                    fulltext = re.sub('([А-Я])\.',r'\1_', result)                  
                    fulltext = re.sub('(\. См\.)|(\.\. См\.\.)',' PUNCT FOREVERALONE', fulltext)
                    sentences = st(fulltext)
                    for sentence in sentences:
                        if 'FOREVERALONE' in sentence and 'Rhomb' not in sentence and 'Tilde' not in sentence:
                            splitter = sentence.split('PUNCT FOREVERALONE')
                            example = splitter[0]
                            example = re.sub('([А-Я])_',r'\1.',example)
                            ref_word = splitter[1]
                            ref_word = ref_word.strip('.')
                            ref_word = ref_word.upper()
                            if 'references' not in entry:
                                entry['references'] = [{'example':example,'reference':ref_word}]
                            else:
                                entry['references'].append({'example':example,'reference':ref_word})
                        elif 'FOREVERALONE' in sentence and ('Rhomb' in sentence or 'Tilde' in sentence):
                            sentence = re.sub('Rhomb','◊',sentence) # Снова вставляем ромбики
                            sentence = re.sub('Tilde','~',sentence) # Снова вставляем тильды
                            splitter = sentence.split('PUNCT FOREVERALONE')
                            example = splitter[0]
                            example = re.sub('([А-Я])_',r'\1.',example)
                            ref_word = splitter[1]
                            ref_word = ref_word.strip('.')
                            ref_word = ref_word.upper()
                            if 'references' not in entry:
                                entry['references'] = [{'example':example,'reference':ref_word}]
                            else:
                                entry['references'].append({'example':example,'reference':ref_word})
                            #bare_sentence = re.sub('''[*.,:;?!()"«»“„_—>']''','',sentence)
                            bare_sentence = re.sub('''[*.,:;?!()"«»“„_—>']''','',example)
                            lemmata = [lemma for lemma in morph.lemmatize(bare_sentence) if lemma not in [' ','\n']]
                            lemmata.append(head.lower())
                            lemmata = ' '.join(lemmata)
                            if 'phrases' not in entry:
                                entry['phrases'] = [{'sentence':f'{example}.','lemmata':lemmata}]
                            else:
                                entry['phrases'].append({'sentence':f'{example}.','lemmata':lemmata})
                        elif 'FOREVERALONE' not in sentence and ('Rhomb' in sentence or 'Tilde' in sentence):
                            bare_sentence = re.sub('''[*.,:;?!()"«»“„_—>']''','',sentence)
                            sentence = re.sub('Rhomb','◊',sentence) # Снова вставляем ромбики
                            sentence = re.sub('Tilde','~',sentence) # Снова вставляем тильды
                            sentence = re.sub('([А-Я])_',r'\1.',sentence)
                            lemmata = [lemma for lemma in morph.lemmatize(bare_sentence) if lemma not in [' ','\n']]
                            lemmata.append(head.lower())
                            lemmata = ' '.join(lemmata)
                            if 'phrases' not in entry:
                                entry['phrases'] = [{'sentence':sentence,'lemmata':lemmata}]
                            else:
                                entry['phrases'].append({'sentence':sentence,'lemmata':lemmata})

                    
                    if 'references' in entry or 'phrases' in entry:
                        entries.append(entry)
                    ##print(heading)
                    entry['headword'] = heading
                    head = re.sub('[0-9]','', re.sub('\u0301','', heading))
                    entry['head'] = head
                    #После этого обнуляем результат
                    entry = {}
                    result = ''
                    #И вносим собственно новую статью, начинающуюся с нашего слова
                    result = result + '\n' + raw_text
                    heading = first_word    

        for entry in entries:
            if 'phrases' in entry:
                hwd = entry['headword']
                phrases = entry['phrases']
                [cur.execute('INSERT INTO for_substrings(headword,lemmata,example,issue) VALUES(?,?,?,?)',[hwd,phr['lemmata'],phr['sentence'],issue])
                    for phr in phrases]
            if 'references' in entry:
                hwd = entry['headword']
                hd = entry['head']
                refs = entry['references']
                [cur.execute('INSERT INTO refers(headword,head,example,link,issue) VALUES(?,?,?,?,?)',[hwd,hd,rf['example'],rf['reference'],issue])
                    for rf in refs]
                
        return entries

    [process(item) for item in files]
    database.commit()

    return {'type':'db_ready','namelist': files}
             
#Далее идут векторные инструменты
    
@applet.route('/make_vectors',methods=['GET','POST'])
def vectorize():
    try:
        modelname = request.json.get('vecmodelname')
        vecsize = int(request.json.get('vecsize'))
        winsize = int(request.json.get('winsize'))
        mincount = int(request.json.get('minimum'))
        epochs = int(request.json.get('epochs'))
        alg = int(request.json.get('alg'))
        if alg == 1:
            algorithm = 'skip-gram'
        elif alg == 0:
            algorithm = 'CBOW'
        global CURRENT_DIR
        current_directory = CURRENT_DIR
        CURRENT_USER = request.json.get('login')
        ##print(CURRENT_USER)
        dbname = request.json.get('dbname')
        source_text = rf'{current_directory}/for_vectors/{CURRENT_USER}/{dbname}.txt'
        magic = gensim.models
        data = magic.word2vec.LineSentence(source_text)
        model = magic.Word2Vec(data,vector_size=vecsize,window=winsize,min_count=mincount,sg=alg,epochs=epochs,workers=3)
        if os.path.isdir(rf'{current_directory}/vecmodels/{CURRENT_USER}'):
            ##print('Такая директория уже существует!')
            pass
        else:
            os.mkdir(rf'{current_directory}/vecmodels/{CURRENT_USER}')
        vecfolder = rf'{current_directory}/vecmodels/{CURRENT_USER}'
        model.save(f'{vecfolder}/{modelname}_{dbname}.model')
        return {'type':'db_ready','chars':f'Размер вектора: {vecsize}, алгоритм: {algorithm} окно: {winsize}, порог: {mincount}, эпохи: {epochs}','modelname':modelname}
    except Exception as err:
        return {'type':'fail','message':f'{err} Проверьте, что указали все характеристики модели и попробуйте снова!'}
    
@applet.route('/choose_vecmodel',methods=['GET','POST'])
def vec_search():
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    models = os.listdir(rf'{current_directory}/vecmodels/{CURRENT_USER}')
    models = [model for model in models if model.endswith('.model')]
    return models

@applet.route('/vecdelete',methods=['GET','POST'])
def del_vectors():
    global CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    current_directory = CURRENT_DIR
    model = request.json.get('modelname')
    os.remove(rf'{current_directory}/vecmodels/{CURRENT_USER}/{model}')
    return {'message':f'Модель {model} успешно удалена!'}

@applet.route('/n_closest',methods=['GET','POST'])
def lookup():
    try:
        global CURRENT_DIR
        current_directory = CURRENT_DIR
        CURRENT_USER = request.json.get('login')
        word_form = request.json.get('word')
        word_form = word_form.lower()
        model = request.json.get('modelname')
        #
        dbname = model.split('_')[-1]
        dbname = re.sub('.model$','',dbname)
        ##print(dbname)
        db = mongo[f'{CURRENT_USER}']
        collection = db[f'{dbname}']
        cardinality = list(collection.find({'cardinality':{'$exists':True}}))[0]
        cardinality = cardinality['cardinality']
        #
        POS = request.json.get('graminfo')
        onlysame = request.json.get('same')
        if POS == 'all':
            pos = 'X'
        else: 
            pos = POS
        word = f'{word_form}_{pos}'
        model = gensim.models.Word2Vec.load(rf'{current_directory}/vecmodels/{CURRENT_USER}/{model}')
        quasi = model.wv.most_similar(positive=word,topn=20)
        quasi = [item for item in quasi if '_' in item[0]]
        header = [f'{word_form.upper()}']
        new_result = []
        #for item in quasi:
        #    ##print(item)
        #    new_item = [item[0].split('_')[0],item[0].split('_')[1],item[1]]
        #    new_result.append(new_item)
        [new_result.append([item[0].split('_')[0],item[0].split('_')[1],item[1]]) for item in quasi]
            
        if onlysame == 'yes':
            result = {'header':header, 'body':[{'wordform':item[0], 'cosine':item[2]} for item in new_result if item[1] == pos]}
        elif onlysame == 'no':
            result = {'header':header, 'body':[{'wordform':item[0], 'cosine':item[2]} for item in new_result]}
    
    except Exception as Err:
        result = {'header':'Слово не найдено!', 'body':[{'wordform':'Вхождения нет','cosine':'попробуйте другое слово!'}]}

    return result

@applet.route('/bigrams',methods=['GET','POST'])
def bigrams():
    global CURRENT_DIR
    current_directory = CURRENT_DIR
    CURRENT_USER = request.json.get('login')
    dbname = request.json.get('dbname')
    stops = request.json.get('stopwords')
    count = int(request.json.get('mincount'))
    necessary=request.json.get('necessary_word')
    source_text = rf'{current_directory}/for_vectors/{CURRENT_USER}/{dbname}.txt'
    text = open(source_text,'r',encoding='UTF-8').read()
    lines = text.split('\n')
    lines = [wt(line) for line in lines if line != '']
    bigram = Phrases(lines,min_count=count,threshold=3,delimiter='~')
    bigram_phraser = Phraser(bigram)
    ##print(bigram_phraser)

    set_phrases = [item for array in [bigram_phraser[line] for line in lines] for item in array]
    if stops == 'yes':
        set_phrases = [item for item in set_phrases if '~' in item]
    elif stops == 'no':
        set_phrases = [item for item in set_phrases if '~' in item and 'PR' not in item and 'CONJ' not in item and 'PART' not in item]
    set_phrases = [re.sub('[_ASPRODVCNJTUM]','',item) for item in set_phrases if '~' in item]
    set_phrases = list(set(set_phrases))
    #set_phrases = [re.sub('~',' ',item) for item in set_phrases]
    #
    set_phrases = [item.split('~') for item in set_phrases]
    if necessary == 'DUMMY':
        set_phrases = [{'frst':item[0],'scnd':item[1]} for item in set_phrases]
    elif necessary != 'DUMMY':
        set_phrases = [{'frst':item[0],'scnd':item[1]} for item in set_phrases if necessary in item]
    #
    if stops=='yes':
        set_phrases = set_phrases
    elif stops=='no':
        set_phrases = [elt for elt in set_phrases if (len(elt['frst'])>1 and len(elt['scnd'])>1)]
    #set_phrases = sorted(set(set_phrases))
    set_phrases = sorted(set_phrases, key=lambda x: x['frst'])

    return {'type':'ngrams_ready','set_phrases': set_phrases, 'error':0}




